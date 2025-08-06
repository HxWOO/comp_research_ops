from openai import OpenAI
import os
import sys
from docx import Document
from dotenv import load_dotenv
import datetime

# --- 전역 설정 ---
load_dotenv()
client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
PROGRESS_FILE = "progress.txt"
COMPANY_LIST_FILE = "kospi_top_100.txt"

def analyze_company(company_name: str, topic: str, job_title: str) -> str:
    """OpenAI API를 사용하여 기업 정보를 분석합니다."""
    print(f"Analyzing '{company_name}' for '{job_title}' on '{topic}'...")
    # AI가 **굵은 글씨**를 사용하도록 프롬프트 수정
    system_prompt = "당신은 저명한 IT 산업 및 기술 전략 분석가입니다. 기업의 역사적 맥락, 현재 상태, 그리고 미래 성장 동력을 깊이 있게 분석하여 신입 지원자에게 통찰력 있는 정보를 제공해야 합니다. 답변 내용 중 가장 중요한 핵심 키워드나 문장은 **굵은 글씨** 처리를 위해 양쪽을 **로 감싸주세요. 다른 마크다운 문법은 절대 사용하지 마세요."
    user_prompts = {
        "기업의 기술적 Legacy 분석": f"'{company_name}'가 창립 이후 겪어온 주요 기술적 변곡점들은 무엇인가요? 과거에 내렸던 중요한 기술적 결정(예: 특정 언어/프레임워크 채택, 아키텍처 설계)들이 현재 시스템에 어떤 '기술 부채(Technical Debt)'나 '유산(Legacy)'으로 남아있는지 분석해 주세요. 그리고 이러한 Legacy를 통해 얻은 교훈은 무엇인지 설명해 주세요.",
        "현재의 주력 사업 및 기술 스택 분석": f"현재 '{company_name}'의 핵심 비즈니스 모델과 주력 서비스는 무엇인가요? 이를 위해 사용하고 있는 최신 기술 스택(언어, 프레임워크, DB, 클라우드, DevOps 등)을 상세히 분석하고, 최근 기술 블로그나 컨퍼런스에서 강조하는 기술 트렌드는 무엇인지 알려주세요.",
        "최근 집중하고 있는 신규 IT 사업 및 투자 분야": f"'{company_name}'가 미래 성장 동력으로 삼고 최근 집중적으로 투자하거나 R&D를 진행하고 있는 신규 IT 사업 분야는 무엇인가요? (예: AI, 블록체인, 메타버스, 신규 플랫폼 등) 관련 자회사 설립, M&A, 대규모 채용 등 구체적인 움직임이 있다면 함께 분석해 주세요.",
        "Legacy와 현재, 그리고 미래로의 기회": f"앞서 분석한 '{company_name}'의 Legacy, 현재 주력 사업, 그리고 미래 신사업 사이에는 어떤 연결고리가 있나요? 회사가 과거의 기술 부채를 해결하고, 현재의 사업을 안정적으로 운영하며, 미래 신사업을 성공시키기 위해 어떤 노력을 하고 있는지 종합적으로 설명해 주세요. 신입 '{job_title}' 개발자가 이 과정에서 어떤 역할을 맡아 기여할 수 있을지, 지원자의 관점에서 기회 포인트를 짚어주세요.",
        "자기소개서 작성을 위한 핵심 전략": f"위 분석 내용을 종합하여, 신입 '{job_title}' 지원자가 자기소개서에 어떤 점을 어필해야 할까요? 회사의 과거(Legacy)에 대한 이해, 현재(State) 기술에 대한 기여 의지, 그리고 미래(Future) 비전에 대한 공감을 동시에 보여줄 수 있는 자기소개서 작성 전략을 구체적인 문장 예시와 함께 3가지 팁으로 제시해 주세요."
    }
    try:
        response = client.chat.completions.create(
            model="gpt-4o", messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompts[topic]}],
            temperature=0.7, max_tokens=2000,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"API call failed: {e}")
        return f"'{topic}' 정보 분석에 실패했습니다."

def generate_report(company_name: str, job_title: str, research_data: dict):
    """분석된 데이터로 Word 보고서를 생성하고, **텍스트**를 굵게 처리합니다."""
    print(f"Generating report for {company_name} ({job_title})...")
    doc = Document()
    doc.add_heading(f"{company_name} '{job_title}' 지원자 맞춤형 심층 분석 보고서", level=0)
    doc.add_paragraph(f"작성일: {datetime.date.today().strftime('%Y년 %m월 %d일')}")
    doc.add_paragraph() # 공백

    for topic, content in research_data.items():
        doc.add_heading(topic, level=1)
        
        # **를 기준으로 텍스트를 분리하여 굵은 글씨 처리
        parts = content.split('**')
        p = doc.add_paragraph() # 새로운 문단 추가
        for i, part in enumerate(parts):
            if i % 2 == 1:  # 홀수 인덱스 (**)로 감싸인 부분)
                p.add_run(part).bold = True
            else:  # 짝수 인덱스 (일반 텍스트)
                p.add_run(part)
        
        doc.add_paragraph() # 섹션 간 간격

    output_dir = "output"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    file_name = f"{output_dir}/{datetime.date.today().strftime('%Y-%m-%d')}_{company_name}_{job_title}_심층_분석_보고서.docx"
    doc.save(file_name)
    print(f"Report saved to '{file_name}'.")

def get_current_rank():
    """progress.txt에서 현재 진행 순위를 읽어옵니다."""
    if not os.path.exists(PROGRESS_FILE):
        return 0
    with open(PROGRESS_FILE, "r") as f:
        return int(f.read().strip())

def update_rank(rank):
    """progress.txt에 다음 진행 순위를 저장합니다."""
    with open(PROGRESS_FILE, "w") as f:
        f.write(str(rank + 1))

if __name__ == "__main__":
    if not os.path.exists(COMPANY_LIST_FILE):
        print(f"Error: '{COMPANY_LIST_FILE}' not found. Please run scraper.py first.")
        sys.exit(1)

    with open(COMPANY_LIST_FILE, "r", encoding="utf-8") as f:
        companies = [line.strip() for line in f.readlines()]

    current_rank = get_current_rank()

    if current_rank >= len(companies):
        print("All companies have been analyzed.")
        sys.exit(0)

    target_company = companies[current_rank]
    print(f"--- Starting analysis for Rank {current_rank + 1}: {target_company} ---")

    job_titles = ["백엔드 개발자", "인프라 엔지니어", "AI 엔지니어"]
    analysis_topics = [
        "기업의 기술적 Legacy 분석", "현재의 주력 사업 및 기술 스택 분석",
        "최근 집중하고 있는 신규 IT 사업 및 투자 분야", "Legacy와 현재, 그리고 미래로의 기회",
        "자기소개서 작성을 위한 핵심 전략"
    ]

    for job in job_titles:
        analysis_results = {}
        for topic in analysis_topics:
            result = analyze_company(target_company, topic, job)
            analysis_results[topic] = result
        generate_report(target_company, job, analysis_results)
    
    update_rank(current_rank)
    print(f"--- Finished analysis for {target_company}. Progress updated. ---")